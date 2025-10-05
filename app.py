import streamlit as st
import pandas as pd
import google.generativeai as genai
import json
from datetime import datetime
from docx import Document
import os
import base64
from dotenv import load_dotenv
def local_css(file_name):
    with open(file_name, encoding='utf-8') as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# نادِ الفانكشن في بداية البرنامج بعد set_page_config
local_css("main.css")

# ================= دالة إضافة الخلفية =================
def add_bg_from_local(image_file):
    with open(image_file, "rb") as f:
        encoded = base64.b64encode(f.read()).decode()
    ext = image_file.split('.')[-1]
    mime_type = f"image/{'svg+xml' if ext == 'svg' else ext}"
    st.markdown(f"""
        <style>
        /* ===== الخلفية الأساسية ===== */
        .stApp {{
            background-image: url("data:{mime_type};base64,{encoded}");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }}

        /* ===== طبقة شفافة فوق الخلفية (تغميق بسيط) ===== */
        .stApp::before {{
            content: "";
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            z-index: 0;
        }}

        /* ===== خلي المحتوى فوق الطبقة ===== */
        .stApp > div:first-child {{
            position: relative;
            z-index: 1;
        }}
        </style>
    """, unsafe_allow_html=True)

# ================= نادِ الخلفية =================
add_bg_from_local("assets/rana.png")
# ===================== إعداد مفتاح Google Generative AI =====================
# Load environment variables from .env file
load_dotenv()
# Get the API key from environment
api_key = os.getenv("GEMINI_API_KEY")
# Configure the Gemini client
genai.configure(api_key=api_key)

# ===================== دوال مساعدة عامة =====================
def _safe_parse_json(raw_text: str):
    """تنظيف وتحويل رد AI إلى JSON صالح"""
    if not isinstance(raw_text, str):
        return None
    txt = raw_text.strip()
    if txt.startswith("```"):
        txt = txt.strip("`")
        if txt.lower().startswith("json"):
            txt = txt[4:].strip()
    try:
        return json.loads(txt)
    except Exception:
        return None

def _to_float(x, default=0.0):
    try:
        if isinstance(x, str):
            x = x.replace("%", "").strip()
        return float(x)
    except Exception:
        return default

# ===================== Benchmarks كلاسيكية + تحليل نصي لتبويبات 3 و 4 =====================
def get_benchmarks_from_ai(category: str):
    """
    تجلب Benchmarks تقديرية للسوق السعودي (CPA, CR, ROAS) وتحوّلها لأرقام آمنة.
    """
    prompt = f"""
    اكتب فقط JSON صالح (بدون أي نص إضافي) لمتوسط مؤشرات السوق السعودي لمجال "{category}".
    استخدم هذه البنية:
    {{
      "CPA": 0.0,
      "CR": 0.0,
      "ROAS": 0.0
    }}
    جميع القيم أرقام (بدون وحدات/رموز).
    واكتب بالعربية لو فيه أسماء حقول إضافية.
    """
    model = genai.GenerativeModel("models/gemini-2.5-flash")
    resp = model.generate_content(prompt)
    data = _safe_parse_json(resp.text or "") or {}
    # تحويل آمن
    return {
        "CPA": _to_float(data.get("CPA", 0)),
        "CR": _to_float(data.get("CR", 0)),
        "ROAS": _to_float(data.get("ROAS", 0)),
    }

def analyze(client, market):
    """
    مقارنة سريعة بين أداء العميل والسوق.
    تتوقع مفاتيح client: CPA(optional), CR, ROAS, NetProfit/ProfitMargin(optional)
    """
    analysis = []
    # CPA (اختياري)
    if "CPA" in client and client["CPA"] is not None:
        if client["CPA"] > market["CPA"]:
            analysis.append(f"🔴 CPA عندك = {client['CPA']:.2f} ريال أعلى من السوق ({market['CPA']:.2f} ريال).")
        else:
            analysis.append(f"🟢 CPA عندك = {client['CPA']:.2f} ريال أفضل من السوق ({market['CPA']:.2f} ريال).")

    # CR
    analysis.append(f"CR = {client['CR']*100:.2f}% مقابل السوق {market['CR']*100:.2f}%.")

    # ROAS
    if client["ROAS"] >= market["ROAS"]:
        analysis.append(f"🟢 ROAS = {client['ROAS']:.2f}x أفضل من السوق ({market['ROAS']:.2f}x).")
    else:
        analysis.append(f"🔴 ROAS = {client['ROAS']:.2f}x أقل من السوق ({market['ROAS']:.2f}x).")

    # ربحية (اختياري)
    if "NetProfit" in client:
        analysis.append(f"صافي الربح/طلب (ريال) = {client['NetProfit']:.2f}")
    elif "ProfitMargin" in client:
        analysis.append(f"هامش الربح/طلب (ريال) = {client['ProfitMargin']:.2f}")

    return "\n".join(analysis)

# ===================== دوال تحليـل الدروب شوبينج (كودك) =====================
def get_ai_analysis(field, CPA, CR, ROAS, orders, visits):
    """جلب Benchmarks السوق + التحليل مباشرة من AI (بالعربية فقط)"""
    today = datetime.today().strftime("%Y-%m-%d")

    # prompt = f"""
    # انت خبير تسويق في السعودية.
    # اكتب الرد بالعربية فقط.
    # التاريخ: {today}
    # المجال: {field}

    # بيانات العميل:
    # - CPA = {CPA:.2f} ريال
    # - CR = {CR*100:.2f}%
    # - ROAS = {ROAS:.2f}x
    # - الأوردرات = {orders}
    # - الزيارات = {visits}

    # اعطني تحليل كامل يتضمن:
    # 1. مؤشرات السوق السعودي اليوم (CPA, CR, ROAS) لمجال {field}.
    # 2. مقارنة بين بيانات العميل والسوق (هل العميل أفضل ✅ أم أضعف ⚠ أم غير منطقي ❌).
    # 3. تحذير إذا كانت بيانات العميل غير منطقية (مثلا CR > 20% أو ROAS > 10x أو زيارات < 100).
    # 4. توصيات عملية لتحسين الأداء.

    # ارجع النتيجة كـ JSON فقط بالصيغة:
    # {{
    #   "MarketBenchmarks": {{"CPA": 0.0, "CR": 0.0, "ROAS": 0.0}},
    #   "Analysis": ["نقطة مقارنة 1", "نقطة مقارنة 2", "نقطة مقارنة 3"],
    #   "Recommendations": ["التوصية 1", "التوصية 2", "التوصية 3"]
    # }}
    # لا تضف أي نص خارج JSON.
    # """
def clean_text_ar(text: str) -> str:
    """تنظيف النص: حذف أي إنجليزي وتصحيح المسافات"""
    import re
    # شيل أي حروف إنجليزية
    text = re.sub(r'[A-Za-z]', '', text)
    # تصحيح المسافات المكررة
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def get_ai_analysis(field, CPA, CR, ROAS, orders, visits):
    """جلب Benchmarks السوق + التحليل مباشرة من AI (بالعربية فقط ومنظم)"""
    today = datetime.today().strftime("%Y-%m-%d")

    prompt = f"""
    انت خبير تسويق في السعودية.
    ✅ مسموح فقط باللغة العربية.
    ❌ ممنوع استخدام أي كلمة أو جملة باللغة الإنجليزية.
    ✅ إذا ذكرت الاختصارات CPA أو CR أو ROAS، يجب أن تكتب بهذا الشكل:
    - تكلفة جذب العميل (CPA)
    - معدل التحويل (CR)
    - عائد الإنفاق الإعلاني (ROAS)
    ✅ اجعل الرد منظم في شكل قائمة مرقمة (1، 2، 3 ...)، بجُمل قصيرة ومباشرة.

    التاريخ: {today}
    المجال: {field}

    بيانات العميل:
    - تكلفة جذب العميل (CPA) = {CPA:.2f} ريال
    - معدل التحويل (CR) = {CR*100:.2f}%
    - عائد الإنفاق الإعلاني (ROAS) = {ROAS:.2f}x
    - الأوردرات = {orders}
    - الزيارات = {visits}

    اعطني تحليل كامل يتضمن:
    1. مؤشرات السوق السعودي الحالية (CPA, CR, ROAS).
    2. مقارنة بين بيانات العميل والسوق (أفضل ✅ – أضعف ⚠ – غير منطقي ❌) ويُعرض بشكل مرقم (1، 2، 3).
    3. تحذيرات إذا كانت البيانات غير منطقية (مثلاً CR > 20% أو ROAS > 10x أو زيارات < 100) وتكون أيضًا مرقمة.
    4. توصيات عملية قصيرة ومباشرة وتكون في شكل قائمة مرقمة.

    النتيجة لازم تكون JSON فقط بالصيغة:
    {{
    "MarketBenchmarks": {{"CPA": 0.0, "CR": 0.0, "ROAS": 0.0}},
    "Analysis": ["1. ...", "2. ...", "3. ..."],
    "Recommendations": ["1. ...", "2. ...", "3. ..."]
    }}
    """
    model = genai.GenerativeModel("models/gemini-2.5-flash")
    response = model.generate_content(prompt)
    data = _safe_parse_json(response.text or "") or {}

    # تطبيع Benchmarks
    if "MarketBenchmarks" in data and isinstance(data["MarketBenchmarks"], dict):
        mb = data["MarketBenchmarks"]
        data["MarketBenchmarks"] = {
            "CPA": _to_float(mb.get("CPA", 0)),
            "CR": _to_float(mb.get("CR", 0)),
            "ROAS": _to_float(mb.get("ROAS", 0)),
        }
    else:
        data["MarketBenchmarks"] = {"CPA": 0.0, "CR": 0.0, "ROAS": 0.0}

    # تنظيف النصوص من أي إنجليزي أو لخبطة
    data["Analysis"] = [clean_text_ar(a) for a in data.get("Analysis", [])]
    data["Recommendations"] = [clean_text_ar(r) for r in data.get("Recommendations", [])]

    return data


    model = genai.GenerativeModel("models/gemini-2.5-flash")
    response = model.generate_content(prompt)
    data = _safe_parse_json(response.text or "") or {}
    # تطبيع الأرقام
    if "MarketBenchmarks" in data and isinstance(data["MarketBenchmarks"], dict):
        mb = data["MarketBenchmarks"]
        data["MarketBenchmarks"] = {
            "CPA": _to_float(mb.get("CPA", 0)),
            "CR": _to_float(mb.get("CR", 0)),
            "ROAS": _to_float(mb.get("ROAS", 0)),
        }
    else:
        data["MarketBenchmarks"] = {"CPA": 0.0, "CR": 0.0, "ROAS": 0.0}
    # تأكد من القوائم
    data.setdefault("Analysis", [])
    data.setdefault("Recommendations", [])
    return data

def export_to_docx(analysis_data, filename="AI_Report.docx"):
    """تحويل تحليل الدروب شوبينج والتوصيات إلى Word"""
    doc = Document()
    doc.add_heading("📑 تقرير السوق (تحليل AI)", 0)

    # Benchmarks السوق
    doc.add_heading("📊 Benchmarks السوق", level=1)
    bm = analysis_data.get("MarketBenchmarks", {})
    doc.add_paragraph(f"CPA (تكلفة الاكتساب) = {bm.get('CPA', 0):.2f} ريال")
    doc.add_paragraph(f"CR (معدل التحويل) = {bm.get('CR', 0)*100:.2f}%")
    doc.add_paragraph(f"ROAS (العائد على الإعلان) = {bm.get('ROAS', 0):.2f}x")

    # التحليل
    doc.add_heading("📊 التحليل", level=1)
    for line in analysis_data.get("Analysis", []):
        doc.add_paragraph(f"• {line}")

    # التوصيات
    doc.add_heading("📌 التوصيات العملية", level=1)
    for rec in analysis_data.get("Recommendations", []):
        doc.add_paragraph(f"- {rec}")

    doc.save(filename)
    return filename

# ====== دالة لتصدير تقرير السوق العام كـ Word (تبويب 5) ======
def export_market_report_to_docx(data, filename="Market_Report.docx"):
    doc = Document()

    # العنوان الرئيسي
    doc.add_heading("📑 تقرير السوق", 0)

    # حجم السوق
    doc.add_heading("📊 حجم السوق (تقديري)", level=1)
    doc.add_paragraph(str(data.get("MarketSize", "-")))

    # معدل النمو السنوي
    doc.add_heading("📈 معدل النمو السنوي (CAGR)", level=1)
    gr = data.get("GrowthRate", "-")
    gr_txt = f"{gr:.2f}%" if isinstance(gr, (int, float)) else str(gr)
    doc.add_paragraph(gr_txt)

    # المنافسين
    doc.add_heading("🏆 أقوى المنافسين في السعودية", level=1)
    comp = data.get("TopCompetitors", []) or []
    for i, c in enumerate(comp, 1):
        doc.add_paragraph(f"{i}. {c}")

    # تحليل SWOT
    doc.add_heading("🔍 تحليل SWOT", level=1)
    sw = data.get("SWOT", {}) or {}

    doc.add_heading("✅ نقاط القوة", level=2)
    for s in sw.get("Strengths", []) or []:
        doc.add_paragraph(f"• {s}")

    doc.add_heading("⚠ نقاط الضعف", level=2)
    for w in sw.get("Weaknesses", []) or []:
        doc.add_paragraph(f"• {w}")

    doc.add_heading("💡 الفرص", level=2)
    for o in sw.get("Opportunities", []) or []:
        doc.add_paragraph(f"• {o}")

    doc.add_heading("🚨 التهديدات", level=2)
    for t in sw.get("Threats", []) or []:
        doc.add_paragraph(f"• {t}")

    # التوصيات
    doc.add_heading("📌 التوصيات", level=1)
    for r in data.get("Recommendations", []) or []:
        doc.add_paragraph(f"• {r}")

    doc.save(filename)
    return filename

# ===================== تحميل البيانات =====================
clients = pd.read_excel("ClientsData_with_SubCategory.xlsx")
df = pd.read_excel("locations_data.xlsx")
countries = df['الدولة'].unique()

# ===================== واجهة Streamlit =====================
st.set_page_config(page_title="تحليل مؤشرات المتاجر", page_icon="", layout="wide")
st.markdown(
    '<h3 class="main-title">نظام تحليل مؤشرات المتاجر بالريال السعودي</h3>',
    unsafe_allow_html=True 
)

# إنشاء التابات
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "BULK عملاء",
    "عميل دروب شوبينج",
    "عميل منتجات خاصه",
    "الأوفلاين بيزنس",
    "السوق"
])

# ========== تبويب 1 ==========
# ===== رفع الملف =====
# st.subheader("⬆ رفع ملف العملاء (Excel)")
# uploaded_file = st.file_uploader("رفع ملف العملاء ", type=["xlsx"])

# if uploaded_file:
#     df_clients = pd.read_excel(uploaded_file)

#     # ===== حساب مؤشرات العميل =====
#     df_clients["CPA"] = df_clients.apply(lambda x: x["الميزانية الإعلانية"]/x["عدد الأوردرات"] if x["عدد الأوردرات"]>0 else 0, axis=1)
#     df_clients["CR"] = df_clients.apply(lambda x: x["عدد الأوردرات"]/x["عدد الزيارات"] if x["عدد الزيارات"]>0 else 0, axis=1)
#     df_clients["ROAS"] = df_clients.apply(lambda x: (x["عدد الأوردرات"]*x["سعر المنتج"])/x["الميزانية الإعلانية"] if x["الميزانية الإعلانية"]>0 else 0, axis=1)

#     # ===== تحليل AI لكل صف =====
#     st.info("⚡ جاري تحليل السوق لكل عميل، يرجى الانتظار...")
#     market_CPAs, market_CRs, market_ROASs, analyses, recs = [], [], [], [], []

#     for idx, row in df_clients.iterrows():
#         field = row["المجال"]
#         ai_result = get_ai_analysis(field, row["CPA"], row["CR"], row["ROAS"], row["عدد الأوردرات"], row["عدد الزيارات"])
#         market_CPAs.append(ai_result["MarketBenchmarks"]["CPA"])
#         market_CRs.append(ai_result["MarketBenchmarks"]["CR"])
#         market_ROASs.append(ai_result["MarketBenchmarks"]["ROAS"])
#         analyses.append(" | ".join(ai_result.get("Analysis",[])))
#         recs.append(" | ".join(ai_result.get("Recommendations",[])))

#     df_clients["Market_CPA"] = market_CPAs
#     df_clients["Market_CR"] = market_CRs
#     df_clients["Market_ROAS"] = market_ROASs
#     df_clients["Analysis"] = analyses
#     df_clients["Recommendations"] = recs

#     st.success("✅ تم اكتمال التحليل لكل العملاء!")

#     # ===== تنزيل Excel =====
#     export_file = f"BULK_Analysis_{datetime.today().strftime('%Y-%m-%d')}.xlsx"
#     df_clients.to_excel(export_file, index=False)
#     with open(export_file, "rb") as f:
#         st.download_button(
#             "⬇ تنزيل ملف التحليل الكامل",
#             data=f,
#             file_name=export_file,
#             mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
#         )

#     st.dataframe(df_clients)
with tab1:
    # ===== رفع الملف =====
    # st.subheader("⬆ رفع ملف العملاء (Excel)")
    uploaded_file = st.file_uploader("⬆ رفع ملف العملاء (Excel)", type=["xlsx"])

    if uploaded_file:
        df_clients = pd.read_excel(uploaded_file)

        # ===== حساب مؤشرات العميل =====
        df_clients["CPA"] = df_clients.apply(lambda x: x["الميزانية الإعلانية"]/x["عدد الأوردرات"] if x["عدد الأوردرات"]>0 else 0, axis=1)
        df_clients["CR"] = df_clients.apply(lambda x: x["عدد الأوردرات"]/x["عدد الزيارات"] if x["عدد الزيارات"]>0 else 0, axis=1)
        df_clients["ROAS"] = df_clients.apply(lambda x: (x["عدد الأوردرات"]*x["سعر المنتج"])/x["الميزانية الإعلانية"] if x["الميزانية الإعلانية"]>0 else 0, axis=1)

        # ===== تحليل AI لكل صف =====
        st.info("⚡ جاري تحليل السوق لكل عميل، يرجى الانتظار...")
        market_CPAs, market_CRs, market_ROASs, analyses, recs = [], [], [], [], []

        for idx, row in df_clients.iterrows():
            field = row["المجال"]
            ai_result = get_ai_analysis(field, row["CPA"], row["CR"], row["ROAS"], row["عدد الأوردرات"], row["عدد الزيارات"])
            market_CPAs.append(ai_result["MarketBenchmarks"]["CPA"])
            market_CRs.append(ai_result["MarketBenchmarks"]["CR"])
            market_ROASs.append(ai_result["MarketBenchmarks"]["ROAS"])
            analyses.append(" | ".join(ai_result.get("Analysis",[])))
            recs.append(" | ".join(ai_result.get("Recommendations",[])))

        df_clients["Market_CPA"] = market_CPAs
        df_clients["Market_CR"] = market_CRs
        df_clients["Market_ROAS"] = market_ROASs
        df_clients["Analysis"] = analyses
        df_clients["Recommendations"] = recs

        st.success("✅ تم اكتمال التحليل لكل العملاء!")

        # ===== تنزيل Excel =====
        export_file = f"BULK_Analysis_{datetime.today().strftime('%Y-%m-%d')}.xlsx"
        df_clients.to_excel(export_file, index=False)
        with open(export_file, "rb") as f:
            st.download_button(
                "⬇ تنزيل ملف التحليل الكامل",
                data=f,
                file_name=export_file,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        st.dataframe(df_clients)

# ========== تبويب 2 (دروب شوبينج بكودك) ==========
# with tab2:
#     st.subheader("➕ إدخال بيانات عميل دروب شوبينج")

#     with st.form("form_ds"):
#         categories = clients["Category"].dropna().unique().tolist()
#         subcategories = clients["SubCategory"].dropna().unique().tolist()

#         category = st.selectbox("اختر المجال (Category)", categories)
#         subcategory = st.selectbox("اختر الفئة الفرعية (SubCategory)", ["None"] + subcategories)

#         price = st.number_input("سعر المنتج (ريال)", min_value=0.0)
#         ad_budget = st.number_input("الميزانية الإعلانية (ريال)", min_value=0.0)
#         orders = st.number_input("عدد الأوردرات", min_value=0.0)
#         visits = st.number_input("عدد الزيارات على المتجر", min_value=0.0)

#         submitted = st.form_submit_button("احسب")

#     if submitted:
#         if orders > visits:
#             st.error("⚠ عدد الأوردرات لا يمكن أن يكون أكبر من عدد الزيارات.")
#         else:
#             # مؤشرات العميل
#             CPA = ad_budget / orders if orders > 0 else 0
#             CR = (orders / visits) if visits > 0 else 0
#             ROAS = (orders * price) / ad_budget if ad_budget > 0 else 0

#             st.markdown("### 🔹 مؤشرات العميل")
#             kpi1, kpi2, kpi3 = st.columns(3)
#             with kpi1:
#                 st.metric("CPA", f"{CPA:.2f} ريال")
#                 st.caption("💡 بدفع كام علشان أجيب عميل جديد")
#             with kpi2:
#                 st.metric("CR", f"{CR*100:.2f}%")
#                 st.caption("💡 نسبة الزوار اللي دخلوا المتجر واشتروا")
#             with kpi3:
#                 st.metric("ROAS", f"{ROAS:.2f}x")
#                 st.caption("💡 بيقيس العائد اللي بيرجعلك مقابل كل ريال إعلان")

#             # المجال النهائي
#             field = f"Dropshipping - {category}"
#             if subcategory != "None":
#                 field += f" - {subcategory}"

#             try:
#                 # تحليل AI
#                 analysis_data = get_ai_analysis(field, CPA, CR, ROAS, orders, visits)

#                 if analysis_data:
#                     st.markdown(f"### 🔹 Benchmarks السوق السعودي ({field})")
#                     bm = analysis_data.get("MarketBenchmarks", {})
#                     bm1, bm2, bm3 = st.columns(3)
#                     with bm1:
#                         st.metric("CPA (متوسط السوق)", f"{bm.get('CPA', 0):.2f} ريال")
#                         st.caption("💡 بدفع كام علشان أجيب عميل جديد")
#                     with bm2:
#                         st.metric("CR (متوسط السوق)", f"{bm.get('CR', 0)*100:.2f}%")
#                         st.caption("💡 نسبة الزوار اللي دخلوا المتجر واشتروا")
#                     with bm3:
#                         st.metric("ROAS (متوسط السوق)", f"{bm.get('ROAS', 0):.2f}x")
#                         st.caption("💡 بيقيس العائد اللي بيرجعلك مقابل كل ريال إعلان")

#                     st.markdown("### 📊 التحليل")
#                     for line in analysis_data.get("Analysis", []):
#                         st.markdown(f"- {line}")

#                     st.markdown("### 📌 التوصيات العملية لتحسين الأداء")
#                     recs = analysis_data.get("Recommendations", [])
#                     if recs:
#                         for rec in recs:
#                             st.markdown(f"- {rec}")
#                     else:
#                         st.warning("⚠ لم تتوفر توصيات لتحسين الأداء.")

#                     # تقرير Word
#                     filename = f"AI_Report_{field}_{datetime.today().strftime('%Y-%m-%d')}.docx"
#                     export_to_docx(analysis_data, filename)
#                     with open(filename, "rb") as f:
#                         st.download_button(
#                             "⬇ تنزيل تقرير السوق (Word)",
#                             data=f,
#                             file_name=filename,
#                             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                         )
#                 else:
#                     st.warning("⚠ لم نتمكن من جلب تحليل السوق من AI.")
#             except Exception as e:
#                 st.error(f"❌ خطأ: {e}")
with tab2:
    # st.markdown('<h3 class="main-title">➕ إدخال بيانات عميل دروب شوبينج</h3>', unsafe_allow_html=True)
    # st.markdown('<div class="form-container">', unsafe_allow_html=True)
    with st.form("form_ds"):
        st.markdown('<h3 class="main-title">➕ إدخال بيانات عميل دروب شوبينج</h3>', unsafe_allow_html=True)
        categories = clients["Category"].dropna().unique().tolist()
        subcategories = clients["SubCategory"].dropna().unique().tolist()

        col1, col2 = st.columns(2)
        with col1:
            category = st.selectbox(
                "اختر المجال (Category)",
                options=categories,
                index=None,
                placeholder="اختر المجال (Category)",
            )
            price = st.number_input("سعر المنتج (ريال)", min_value=0.0)
            orders = st.number_input("عدد الأوردرات", min_value=0.0)
        with col2:
            subcategory = st.selectbox(
                "اختر الفئة الفرعية (SubCategory)",
                options=subcategories,
                index=None,
                placeholder="اختر الفئة الفرعية (SubCategory)",
            )
            ad_budget = st.number_input("الميزانية الإعلانية (ريال)", min_value=0.0)
            visits = st.number_input("عدد الزيارات على المتجر", min_value=0.0)

        submitted = st.form_submit_button("احسب")

    st.markdown('</div>', unsafe_allow_html=True)

# ========== تبويب 3 (منتجات خاصة - مع اختيار المجال من الداتا) ==========
with tab3:
    with st.form("form_private"):
        st.subheader("➕ إدخال بيانات عميل منتجات خاصة")
        # ↓↓↓ جلب القوائم من ملف العملاء
        categories = clients["Category"].dropna().unique().tolist() if "Category" in clients.columns else []
        subcategories = clients["SubCategory"].dropna().unique().tolist() if "SubCategory" in clients.columns else []

        ccol1, ccol2 = st.columns(2)
        with ccol1:
            category = st.selectbox(
                "اختر المجال (Category)",
                options=categories,
                index=None,
                placeholder="اختر المجال (Category)",
                key="pp_category",
            )
        with ccol2:
            subcategory = st.selectbox(
                "اختر الفئة الفرعية (SubCategory)",
                options=subcategories,
                index=None,
                placeholder="اختر الفئة الفرعية (SubCategory)",
                key="pp_subcategory",
            )

        # ↓↓↓ مدخلات الحساب (صفوف، كل صف يحتوي مدخلين)
        r1c1, r1c2 = st.columns(2)
        with r1c1:
            price = st.number_input("سعر المنتج (ريال)", min_value=0.0, key="pp_price")
        with r1c2:
            ad_budget = st.number_input("الميزانية الإعلانية (ريال)", min_value=0.0, key="pp_budget")

        r2c1, r2c2 = st.columns(2)
        with r2c1:
            orders = st.number_input("عدد الأوردرات", min_value=0.0, key="pp_orders")
        with r2c2:
            visits = st.number_input("عدد الزيارات على المتجر", min_value=0.0, key="pp_visits")

        submitted = st.form_submit_button("احسب", key="pp_submit")

    if submitted:
        if orders > visits:
            st.error("⚠ عدد الأوردرات لا يمكن أن يكون أكبر من عدد الزيارات.")
        else:
            # حساب مؤشرات العميل
            CPA = ad_budget / orders if orders > 0 else 0
            CR  = (orders / visits) if visits > 0 else 0
            ROAS = (orders * price) / ad_budget if ad_budget > 0 else 0

            # عرض النتائج
            st.markdown("""<div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
            <h3 style="color: white; text-align: center;">🔹 مؤشرات العميل</h3>
            </div>""",     
            unsafe_allow_html=True )
            k1, k2, k3 = st.columns(3)
            k1.metric("CPA", f"{CPA:.2f} ريال")
            k2.metric("CR", f"{CR*100:.2f}%")
            k3.metric("ROAS", f"{ROAS:.2f}x")

            # المجال لتحليل الذكاء الاصطناعي
            field = "Private Products"
            if category:
                field += f" - {category}"
            if subcategory:
                field += f" - {subcategory}"

            try:
                analysis_data = get_ai_analysis(field, CPA, CR, ROAS, orders, visits)
                if analysis_data:
                    # Benchmarks السوق
                    st.markdown(f"""<div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
                    <h3 style="color: white; text-align: center;">
                    🔹 Benchmarks السوق السعودي ({field})
                    </h3></div>""",     
                    unsafe_allow_html=True )
                    # st.markdown(f"### 🔹 Benchmarks السوق السعودي ({field})")
                    bm = analysis_data.get("MarketBenchmarks", {})
                    c1, c2, c3 = st.columns(3)
                    c1.metric("CPA (متوسط السوق)", f"{bm.get('CPA', 0):.2f} ريال")
                    c2.metric("CR (متوسط السوق)", f"{bm.get('CR', 0)*100:.2f}%")
                    c3.metric("ROAS (متوسط السوق)", f"{bm.get('ROAS', 0):.2f}x")

                    # التحليل
                    # st.markdown("### 📊 التحليل")
                    st.markdown("""<div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
                    <h3 style="color: white; text-align: center;">📊 التحليل</h3>
                    </div>""",     
                    unsafe_allow_html=True )
                    for line in analysis_data.get("Analysis", []):
                        # st.markdown(f"- {line}")
                        st.markdown(f"""
                        <div style="
                            position: relative;
                            background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
                            backdrop-filter: blur(10px);
                            border: 1px solid rgba(255,255,255,0.2);
                            box-shadow: 0 2px 8px rgba(0,0,0,0.3);
                            border-radius: 10px;
                            padding: 10px 14px;
                            margin: 8px 0;
                        ">
                            <p style="
                                color: #f0f0f0;
                                font-size: 0.95rem;
                                font-family: 'Cairo', sans-serif;
                                line-height: 1.6;
                                text-align: right;
                                direction: rtl;
                            ">🔹 {line}</p>
                        </div>
                        """, unsafe_allow_html=True)

                    # التوصيات
                    # st.markdown("### 📌 التوصيات العملية لتحسين الأداء")
                    st.markdown("""<div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
                    <h3 style="color: white; text-align: center;">📌 التوصيات العملية لتحسين الأداء</h3>
                    </div>""",     
                    unsafe_allow_html=True )
                    for rec in analysis_data.get("Recommendations", []):
                        # st.markdown(f"- {rec}")
                        st.markdown(f"""
                        <div style="
                            position: relative;
                            background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
                            backdrop-filter: blur(10px);
                            border: 1px solid rgba(255,255,255,0.2);
                            box-shadow: 0 2px 8px rgba(0,0,0,0.3);
                            border-radius: 10px;
                            padding: 10px 14px;
                            margin: 8px 0;
                        ">
                            <p style="
                                color: #f0f0f0;
                                font-size: 0.95rem;
                                font-family: 'Cairo', sans-serif;
                                line-height: 1.6;
                                text-align: right;
                                direction: rtl;
                            ">🔹 {rec}</p>
                        </div>
                        """, unsafe_allow_html=True)

                    # تقرير Word
                    filename = f"AI_Report_{field}_{datetime.today().strftime('%Y-%m-%d')}.docx"
                    export_to_docx(analysis_data, filename)
                    with open(filename, "rb") as f:
                        st.download_button(
                            "⬇ تنزيل تقرير السوق (Word)",
                            data=f,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.warning("⚠ لم نتمكن من جلب تحليل السوق من AI.")
            except Exception as e:
                st.error(f"❌ خطأ: {e}")

# ========== تبويب 4 (بزنس الأوفلاين - بنفس أسلوب التبويبات الأخرى) ==========
with tab4:
    with st.form("form_offline"):
        st.subheader("➕ إدخال بيانات الأوفلاين بيزنس ")
        # اختيار المجال والفئة الفرعية
        categories = clients["Category"].dropna().unique().tolist() if "Category" in clients.columns else []
        subcategories = clients["SubCategory"].dropna().unique().tolist() if "SubCategory" in clients.columns else []

        ocol1, ocol2 = st.columns(2)
        with ocol1:
            category = st.selectbox(
                "اختر المجال (Category)",
                options=categories,
                index=None,
                placeholder="اختر المجال (Category)",
                key="off_category",
            )
        with ocol2:
            subcategory = st.selectbox(
                "اختر الفئة الفرعية (SubCategory)",
                options=subcategories,
                index=None,
                placeholder="اختر الفئة الفرعية (SubCategory)",
                key="off_subcategory",
            )

        # مدخلات موحدة (صفوف، كل صف يحتوي مدخلين)
        r1c1, r1c2 = st.columns(2)
        with r1c1:
            price = st.number_input("سعر المنتج (ريال)", min_value=0.0, key="off_price")
        with r1c2:
            ad_budget = st.number_input("الميزانية الإعلانية (ريال)", min_value=0.0, key="off_budget")

        r2c1, r2c2 = st.columns(2)
        with r2c1:
            orders = st.number_input("عدد الأوردرات", min_value=0.0, key="off_orders")
        with r2c2:
            visits = st.number_input("عدد الزيارات على المتجر", min_value=0.0, key="off_visits")

        submitted = st.form_submit_button("احسب", key="off_submit")

    if submitted:
        if orders > visits:
            st.error("⚠ عدد الأوردرات لا يمكن أن يكون أكبر من عدد الزيارات.")
        else:
            # حساب المؤشرات
            CPA = ad_budget / orders if orders > 0 else 0
            CR  = (orders / visits) if visits > 0 else 0
            ROAS = (orders * price) / ad_budget if ad_budget > 0 else 0

            # عرض النتائج
            st.markdown("### 🔹 مؤشرات العميل")
            o1, o2, o3 = st.columns(3)
            o1.metric("CPA", f"{CPA:.2f} ريال")
            o2.metric("CR", f"{CR*100:.2f}%")
            o3.metric("ROAS", f"{ROAS:.2f}x")

            # تحديد المجال للتحليل
            field = "Offline Business"
            if category:
                field += f" - {category}"
            if subcategory:
                field += f" - {subcategory}"

            try:
                # تحليل AI
                analysis_data = get_ai_analysis(field, CPA, CR, ROAS, orders, visits)

                if analysis_data:
                    # Benchmarks السوق
                    st.markdown(f"### 🔹 Benchmarks السوق السعودي ({field})")
                    bm = analysis_data.get("MarketBenchmarks", {})
                    b1, b2, b3 = st.columns(3)
                    b1.metric("CPA (متوسط السوق)", f"{bm.get('CPA', 0):.2f} ريال")
                    b2.metric("CR (متوسط السوق)", f"{bm.get('CR', 0)*100:.2f}%")
                    b3.metric("ROAS (متوسط السوق)", f"{bm.get('ROAS', 0):.2f}x")

                    # التحليل
                    st.markdown("### 📊 التحليل")
                    for line in analysis_data.get("Analysis", []):
                        st.markdown(f"- {line}")

                    # التوصيات
                    st.markdown("### 📌 التوصيات العملية لتحسين الأداء")
                    for rec in analysis_data.get("Recommendations", []):
                        st.markdown(f"- {rec}")

                    # تقرير Word
                    filename = f"AI_Report_{field}_{datetime.today().strftime('%Y-%m-%d')}.docx"
                    export_to_docx(analysis_data, filename)
                    with open(filename, "rb") as f:
                        st.download_button(
                            "⬇ تنزيل تقرير السوق (Word)",
                            data=f,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.warning("⚠ لم نتمكن من جلب تحليل السوق من AI.")
            except Exception as e:
                st.error(f"❌ خطأ: {e}")


# ========== تبويب 5 ==========
# with tab5:
#     st.subheader("📈 تحليل السوق السعودي ")

#     colA, colB, colC, colD = st.columns(4)
#     with colA:
#         store_market = st.selectbox("اختر المتجر", clients["StoreName"].unique(), key="store_market")
#     with colB:
#         btype = st.selectbox("اختر نوع الفئة", ["B2C", "B2B"], key="btype_market")
#     with colC:
#         country = st.selectbox("اختر الدولة", countries, key="country_market")
#     with colD:
#         cities = df[df['الدولة'] == country]['المنطقة'].unique()
#         cities_options = ["None"] + list(cities)
#         cities_selected = st.multiselect("اختر المدن", cities_options, key="cities_market")

#     if not cities_selected or "None" in cities_selected:
#         selected_text = f"كل مدن {country}"
#         cities_selected = list(cities)
#     else:
#         selected_text = ", ".join(cities_selected)

#     if st.button(" تحليل السوق بالذكاء الاصطناعي"):
#         try:
#             row_m = clients[clients["StoreName"] == store_market].iloc[0]
#             category_market = row_m.get("SubCategory") if pd.notna(row_m.get("SubCategory")) else row_m["Category"]

#             # prompt_market = f"""
#             # اكتب بالعربية فقط.
#             # اعطني تقرير تحليلي موثق عن السوق السعودي في مجال "{category_market}" للفئة "{btype}" 
#             # في دولة {country} ومدن {selected_text}.

#             # يجب أن يتضمن التقرير:
#             # - حجم السوق (بالريال السعودي أو عدد العملاء).
#             # - معدل النمو السنوي (CAGR).
#             # - أقوى 3 منافسين حقيقيين.
#             # - تحليل SWOT.
#             # - 3 توصيات عملية.

#             # ارجع النتيجة في JSON فقط بالصيغة:
#             # {{
#             #   "MarketSize": "...",
#             #   "GrowthRate": 0.0,
#             #   "TopCompetitors": ["...", "...", "..."],
#             #   "SWOT": {{
#             #     "Strengths": ["..."],
#             #     "Weaknesses": ["..."],
#             #     "Opportunities": ["..."],
#             #     "Threats": ["..."]
#             #   }},
#             #   "Recommendations": ["...", "...", "..."]
#             # }}
#             # لا تكتب أي نص خارج JSON.
#             # """
# #             prompt_market = f"""
# # انت باحث تسويق متخصص في السعودية. 
# # ❌ لا تكتب أي كلمات إنجليزية. 
# # ✅ النتيجة لازم تكون بالعربية المبسطة فقط.
# # ✅ اجعل كل جزء في نقاط قصيرة ومنظمة (Bullet Points).

# # اعطني تقرير عن السوق السعودي في مجال "{category_market}" للفئة "{btype}" 
# # في دولة {country} ومدن {selected_text}.

# # يجب أن يتضمن التقرير:
# # - حجم السوق (بالريال السعودي أو عدد العملاء).
# # - معدل النمو السنوي (CAGR).
# # - أقوى 3 منافسين حقيقيين.
# # - تحليل SWOT (نقاط القوة، الضعف، الفرص، التهديدات).
# # - 3 توصيات عملية واضحة.

# # النتيجة لازم تكون JSON فقط بالصيغة:
# # {{
# #   "MarketSize": "...",
# #   "GrowthRate": 0.0,
# #   "TopCompetitors": ["...", "...", "..."],
# #   "SWOT": {{
# #     "Strengths": ["..."],
# #     "Weaknesses": ["..."],
# #     "Opportunities": ["..."],
# #     "Threats": ["..."]
# #   }},
# #   "Recommendations": ["...", "...", "..."]
# # }}
# # """
#             prompt_market = f"""
# انت باحث تسويق متخصص في السعودية. 
# ✅ مسموح فقط باللغة العربية المبسطة.
# ❌ ممنوع استخدام أي كلمة أو جملة باللغة الإنجليزية.
# ✅ لو لازم تذكر مصطلحات عالمية، اكتبها بالعربية متبوعة بالاختصار بين أقواس، مثل:
#    - معدل النمو السنوي المركب (CAGR)
# ✅ اجعل كل جزء من التقرير في شكل قائمة مرقمة (1، 2، 3 ...).

# اعطني تقرير عن السوق السعودي في مجال "{category_market}" للفئة "{btype}" 
# في دولة {country} ومدن {selected_text}.

# يجب أن يتضمن التقرير:
# 1. حجم السوق (بالريال السعودي أو عدد العملاء).
# 2. معدل النمو السنوي المركب (CAGR).
# 3. أقوى 3 منافسين حقيقيين.
# 4. تحليل SWOT (نقاط القوة، الضعف، الفرص، التهديدات) – كل قسم مرقم.
# 5. 3 توصيات عملية واضحة ومباشرة.

# النتيجة لازم تكون JSON فقط بالصيغة:
# {{
#   "MarketSize": "...",
#   "GrowthRate": 0.0,
#   "TopCompetitors": ["1. ...", "2. ...", "3. ..."],
#   "SWOT": {{
#     "Strengths": ["1. ...", "2. ..."],
#     "Weaknesses": ["1. ...", "2. ..."],
#     "Opportunities": ["1. ...", "2. ..."],
#     "Threats": ["1. ...", "2. ..."]
#   }},
#   "Recommendations": ["1. ...", "2. ...", "3. ..."]
# }}
# """
            

#             model = genai.GenerativeModel("models/gemini-2.5-flash")
#             response = model.generate_content(prompt_market)
#             data = _safe_parse_json(response.text or "") or {}

#             st.markdown("### 📊 حجم السوق (تقديري)")
#             st.write(str(data.get("MarketSize", "-")))

#             st.markdown("### 📈 معدل النمو السنوي (CAGR)")
#             gr = data.get("GrowthRate", "-")
#             gr_txt = f"{_to_float(gr):.2f}%"
#             st.write(gr_txt)

#             st.markdown("### 🏆 أقوى المنافسين في السعودية")
#             comp = data.get("TopCompetitors", []) or []
#             if comp:
#                 comp_df = pd.DataFrame({"#": list(range(1, len(comp)+1)), "المنافس": comp})
#                 st.table(comp_df)
#             else:
#                 st.warning("لم يتم العثور على منافسين.")

#             st.markdown("### 🔍 تحليل SWOT")
#             sw = data.get("SWOT", {}) or {}
#             st.write("✅ نقاط القوة")
#             st.write("\n".join([f"• {x}" for x in sw.get("Strengths", [])]) or "-")
#             st.write("⚠ نقاط الضعف")
#             st.write("\n".join([f"• {x}" for x in sw.get("Weaknesses", [])]) or "-")
#             st.write("💡 الفرص")
#             st.write("\n".join([f"• {x}" for x in sw.get("Opportunities", [])]) or "-")
#             st.write("🚨 التهديدات")
#             st.write("\n".join([f"• {x}" for x in sw.get("Threats", [])]) or "-")

#             st.markdown("### 📌 التوصيات")
#             recs = data.get("Recommendations", []) or []
#             st.write("\n".join([f"• {r}" for r in recs]) or "-")

#             # تنزيل Word
#             file_suffix = f"{country}{''.join(cities_selected) if cities_selected else 'عام'}"
#             filename = f"Market_Report_{store_market}{btype}{file_suffix}.docx"
#             export_market_report_to_docx(data, filename)
#             with open(filename, "rb") as f:
#                 st.download_button(
#                     "⬇ تنزيل تقرير السوق (Word)",
#                     data=f,
#                     file_name=filename,
#                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                 )

#         except Exception as e:
#             st.error(f"❌ خطأ في تحليل السوق: {e}")
# with tab5:
#     with st.form("form_market_analysis"):
#         st.subheader("📈 تحليل السوق السعودي ")

#         # صف 1: مدخلان
#         r1c1, r1c2 = st.columns(2)
#         with r1c1:
#             store_market = st.selectbox(
#                 "اختر المتجر",
#                 options=clients["StoreName"].unique(),
#                 index=None,
#                 placeholder="اختر المتجر",
#                 key="store_market",
#             )
#         with r1c2:
#             btype = st.selectbox(
#                 "اختر نوع الفئة",
#                 options=["B2C", "B2B"],
#                 index=None,
#                 placeholder="اختر نوع الفئة",
#                 key="btype_market",
#             )

#         # صف 2: مدخلان
#         r2c1, r2c2 = st.columns(2)
#         with r2c1:
#             country = st.selectbox(
#                 "اختر الدولة",
#                 options=countries,
#                 index=None,
#                 placeholder="اختر الدولة",
#                 key="country_market",
#             )
#         with r2c2:
#             cities = df[df['الدولة'] == country]['المنطقة'].unique()
#             cities_options = ["None"] + list(cities)
#             cities_selected = st.multiselect("اختر المدن", cities_options, key="cities_market")

#         if not cities_selected or "None" in cities_selected:
#             selected_text = f"كل مدن {country}"
#             cities_selected = list(cities)
#         else:
#             selected_text = ", ".join(cities_selected)

#         if st.button(" تحليل السوق بالذكاء الاصطناعي"):
#             try:
#                 row_m = clients[clients["StoreName"] == store_market].iloc[0]
#                 category_market = row_m.get("SubCategory") if pd.notna(row_m.get("SubCategory")) else row_m["Category"]

#                 # ==== Prompt مرتب ====
#                 prompt_market = f"""
#                 انت باحث تسويق متخصص في السعودية. 
#                 ✅ مسموح فقط باللغة العربية المبسطة.
#                 ❌ ممنوع استخدام أي كلمة أو جملة باللغة الإنجليزية.
#                 ✅ لو لازم تذكر مصطلحات عالمية، اكتبها بالعربية متبوعة بالاختصار بين أقواس، مثل:
#                 - معدل النمو السنوي المركب (CAGR)
#                 ✅ اجعل كل جزء من التقرير في شكل قائمة مرقمة (1. ... 2. ... 3. ...).
#                 ✅ كل نقطة لازم تكون جملة قصيرة ومباشرة (سطر واحد فقط).

#                 اعطني تقرير عن السوق السعودي في مجال "{category_market}" للفئة "{btype}" 
#                 في دولة {country} ومدن {selected_text}.

#                 يجب أن يتضمن التقرير:
#                 1. حجم السوق (بالريال السعودي أو عدد العملاء).
#                 2. معدل النمو السنوي المركب (CAGR).
#                 3. أقوى 3 منافسين حقيقيين.
#                 4. تحليل SWOT (نقاط القوة، الضعف، الفرص، التهديدات) – كل قسم مرقم.
#                 5. 3 توصيات عملية واضحة ومباشرة.

#                 النتيجة لازم تكون JSON فقط بالصيغة:
#                 {{
#                 "MarketSize": "...",
#                 "GrowthRate": 0.0,
#                 "TopCompetitors": [". ...", ". ...", ". ..."],
#                 "SWOT": {{
#                     "Strengths": [" ...", " ..."],
#                     "Weaknesses": [" ...", " ..."],
#                     "Opportunities": ["...", "..."],
#                     "Threats": [" ...", " ..."]
#                 }},
#                 "Recommendations": ["...", "...", "..."]
#                 }}
#                 """
#                 model = genai.GenerativeModel("models/gemini-2.5-flash")
#                 response = model.generate_content(prompt_market)
#                 data = _safe_parse_json(response.text or "") or {}

#                 # ==== عرض النتائج في Streamlit ====
#                 # st.markdown("### 📊 حجم السوق (تقديري)")
#                 st.markdown("""
#                 <div style="position: relative; background: rgba(0, 0, 0, 0.5);
#                 backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> 📊 حجم السوق (تقديري)</h3>
#                 </div>
#                 """, unsafe_allow_html=True)
#                 st.write(str(data.get("MarketSize", "-")))

#                 # st.markdown("### 📈 معدل النمو السنوي (CAGR)")
#                 st.markdown("""
#                 <div style="position: relative; background: rgba(0, 0, 0, 0.5);
#                 backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> 📈 معدل النمو السنوي (CAGR)</h3>
#                 </div>
#                 """, unsafe_allow_html=True)
#                 gr = data.get("GrowthRate", "-")
#                 gr_txt = f"{_to_float(gr):.2f}%"
#                 st.write(gr_txt)

#                 # st.markdown("###  أقوى المنافسين في السعودية")
#                 st.markdown(""" 
#                 <div style="position: relative; background: rgba(0, 0, 0, 0.5);
#                 backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> أقوى المنافسين في السعودية</h3>
#                 </div>
#                 """, unsafe_allow_html=True)
#                 comp = data.get("TopCompetitors", []) or []
#                 if comp:
#                     for c in comp:
#                         st.markdown(f"- {c}")
#                 else:
#                     st.warning("لم يتم العثور على منافسين.")

#                 # st.markdown("### تحليل SWOT")
#                 st.markdown("""
#                 <div style="position: relative; background: rgba(0, 0, 0, 0.5);
#                 backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> تحليل SWOT</h3>
#                 </div>
#                 """, unsafe_allow_html=True)
#                 sw = data.get("SWOT", {}) or {}

#                 st.write("✅ نقاط القوة")
#                 for i, x in enumerate(sw.get("Strengths", []), 1):
#                     # st.markdown(f"{i}. {x}")
#                     st.markdown(f"""
#                     <div style="
#                         position: relative;
#                         background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
#                         backdrop-filter: blur(10px);
#                         border: 1px solid rgba(255,255,255,0.2);
#                         box-shadow: 0 2px 8px rgba(0,0,0,0.3);
#                         border-radius: 10px;
#                         padding: 10px 14px;
#                         margin: 8px 0;
#                     ">
#                         <p style="
#                             color: #f0f0f0;
#                             font-size: 0.95rem;
#                             font-family: 'Cairo', sans-serif;
#                             line-height: 1.6;
#                             text-align: right;
#                             direction: rtl;
#                         ">🔹 {i}. {x}</p>
#                     </div>
#                     """, unsafe_allow_html=True)

#                 st.write("⚠ نقاط الضعف")
#                 for i, x in enumerate(sw.get("Weaknesses", []), 1):
#                     # st.markdown(f"{i}. {x}")
#                     st.markdown(f"""
#                     <div style="
#                         position: relative;
#                         background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
#                         backdrop-filter: blur(10px);
#                         border: 1px solid rgba(255,255,255,0.2);
#                         box-shadow: 0 2px 8px rgba(0,0,0,0.3);
#                         border-radius: 10px;
#                         padding: 10px 14px;
#                         margin: 8px 0;
#                     ">
#                         <p style="
#                             color: #f0f0f0;
#                             font-size: 0.95rem;
#                             font-family: 'Cairo', sans-serif;
#                             line-height: 1.6;
#                             text-align: right;
#                             direction: rtl;
#                         ">🔹 {i}. {x}</p>
#                     </div>
#                     """, unsafe_allow_html=True)

#                 st.write("💡 الفرص")
#                 for i, x in enumerate(sw.get("Opportunities", []), 1):
#                     # st.markdown(f"{i}. {x}")
#                     st.markdown(f"""
#                     <div style="
#                         position: relative;
#                         background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
#                         backdrop-filter: blur(10px);
#                         border: 1px solid rgba(255,255,255,0.2);
#                         box-shadow: 0 2px 8px rgba(0,0,0,0.3);
#                         border-radius: 10px;
#                         padding: 10px 14px;
#                         margin: 8px 0;
#                     ">
#                         <p style="
#                             color: #f0f0f0;
#                             font-size: 0.95rem;
#                             font-family: 'Cairo', sans-serif;
#                             line-height: 1.6;
#                             text-align: right;
#                             direction: rtl;
#                         ">🔹 {i}. {x}</p>
#                     </div>
#                     """, unsafe_allow_html=True)

#                 st.write("🚨 التهديدات")
#                 for i, x in enumerate(sw.get("Threats", []), 1):
#                     # st.markdown(f"{i}. {x}")
#                     st.markdown(f"""
#                             <div style="
#                                 position: relative;
#                                 background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
#                                 backdrop-filter: blur(10px);
#                                 border: 1px solid rgba(255,255,255,0.2);
#                                 box-shadow: 0 2px 8px rgba(0,0,0,0.3);
#                                 border-radius: 10px;
#                                 padding: 10px 14px;
#                                 margin: 8px 0;
#                             ">
#                                 <p style="
#                                     color: #f0f0f0;
#                                     font-size: 0.95rem;
#                                     font-family: 'Cairo', sans-serif;
#                                     line-height: 1.6;
#                                     text-align: right;
#                                     direction: rtl;
#                                 ">🔹 {i}. {x}</p>
#                             </div>
#                             """, unsafe_allow_html=True)

#                 # st.markdown("### 📌 التوصيات")
#                 st.markdown("""
#                 <div style="position: relative; background: rgba(0, 0, 0, 0.5);
#                 backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;">📌 التوصيات</h3>
#                 </div>
#                 """, unsafe_allow_html=True)
#                 recs = data.get("Recommendations", []) or []
#                 for i, r in enumerate(recs, 1):
#                     # st.markdown(f"{i}. {r}")
#                     st.markdown(f"""
#                             <div style="
#                                 position: relative;
#                                 background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
#                                 backdrop-filter: blur(10px);
#                                 border: 1px solid rgba(255,255,255,0.2);
#                                 box-shadow: 0 2px 8px rgba(0,0,0,0.3);
#                                 border-radius: 10px;
#                                 padding: 10px 14px;
#                                 margin: 8px 0;
#                             ">
#                                 <p style="
#                                     color: #f0f0f0;
#                                     font-size: 0.95rem;
#                                     font-family: 'Cairo', sans-serif;
#                                     line-height: 1.6;
#                                     text-align: right;
#                                     direction: rtl;
#                                 ">🔹 {i}. {r}</p>
#                             </div>
#                             """, unsafe_allow_html=True)

#                 # ==== تنزيل Word ====
#                 file_suffix = f"{country}{''.join(cities_selected) if cities_selected else 'عام'}"
#                 filename = f"Market_Report_{store_market}{btype}{file_suffix}.docx"
#                 export_market_report_to_docx(data, filename)
#                 with open(filename, "rb") as f:
#                     st.download_button(
#                         "⬇ تنزيل تقرير السوق (Word)",
#                         data=f,
#                         file_name=filename,
#                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                     )

#             except Exception as e:
#                 st.error(f"❌ خطأ في تحليل السوق: {e}")

# ========== تبويب تحليل السوق السعودي مع نموذج إدخال مثل نموذج الأوفلاين ==========
# with tab5:
#     with st.form("form_market_analysis"):
#         st.subheader("📈 تحليل السوق السعودي")

#         # صف 1: مدخلان
#         r1c1, r1c2 = st.columns(2)
#         with r1c1:
#             store_market = st.selectbox(
#                 "اختر المتجر",
#                 options=clients["StoreName"].unique(),
#                 index=None,
#                 placeholder="اختر المتجر",
#                 key="store_market",
#             )
#         with r1c2:
#             btype = st.selectbox(
#                 "اختر نوع الفئة",
#                 options=["B2C", "B2B"],
#                 index=None,
#                 placeholder="اختر نوع الفئة",
#                 key="btype_market",
#             )

#         # صف 2: مدخلان
#         r2c1, r2c2 = st.columns(2)
#         with r2c1:
#             country = st.selectbox(
#                 "اختر الدولة",
#                 options=countries,
#                 index=None,
#                 placeholder="اختر الدولة",
#                 key="country_market",
#             )
#         with r2c2:
#             cities = df[df['الدولة'] == country]['المنطقة'].unique()
#             cities_options = ["None"] + list(cities)
#             cities_selected = st.multiselect("اختر المدن", cities_options, key="cities_market")

#         if not cities_selected or "None" in cities_selected:
#             selected_text = f"كل مدن {country}"
#             cities_selected = list(cities)
#         else:
#             selected_text = ", ".join(cities_selected)

#         # زر لتحليل السوق بالذكاء الاصطناعي
#         submitted = st.form_submit_button("تحليل السوق بالذكاء الاصطناعي")
        
#         if submitted:
#             try:
#                 row_m = clients[clients["StoreName"] == store_market].iloc[0]
#                 category_market = row_m.get("SubCategory") if pd.notna(row_m.get("SubCategory")) else row_m["Category"]

#                 # إنشاء Prompt مرتب للتحليل
#                 prompt_market = f"""
#                 انت باحث تسويق متخصص في السعودية. 
#                 ✅ مسموح فقط باللغة العربية المبسطة.
#                 ❌ ممنوع استخدام أي كلمة أو جملة باللغة الإنجليزية.
#                 ✅ لو لازم تذكر مصطلحات عالمية، اكتبها بالعربية متبوعة بالاختصار بين أقواس، مثل:
#                 - معدل النمو السنوي المركب (CAGR)
#                 ✅ اجعل كل جزء من التقرير في شكل قائمة مرقمة (1. ... 2. ... 3. ...).
#                 ✅ كل نقطة لازم تكون جملة قصيرة ومباشرة (سطر واحد فقط).

#                 اعطني تقرير عن السوق السعودي في مجال "{category_market}" للفئة "{btype}" 
#                 في دولة {country} ومدن {selected_text}.

#                 يجب أن يتضمن التقرير:
#                 1. حجم السوق (بالريال السعودي أو عدد العملاء).
#                 2. معدل النمو السنوي المركب (CAGR).
#                 3. أقوى 3 منافسين حقيقيين.
#                 4. تحليل SWOT (نقاط القوة، الضعف، الفرص، التهديدات) – كل قسم مرقم.
#                 5. 3 توصيات عملية واضحة ومباشرة.

#                 النتيجة لازم تكون JSON فقط بالصيغة:
#                 {{
#                 "MarketSize": "...",
#                 "GrowthRate": 0.0,
#                 "TopCompetitors": [". ...", ". ...", ". ..."],
#                 "SWOT": {{
#                     "Strengths": [" ...", " ..."],
#                     "Weaknesses": [" ...", " ..."],
#                     "Opportunities": ["...", "..."],
#                     "Threats": [" ...", " ..."]
#                 }},
#                 "Recommendations": ["...", "...", "..."]
#                 }}
#                 """

#                 model = genai.GenerativeModel("models/gemini-2.5-flash")
#                 response = model.generate_content(prompt_market)
#                 data = _safe_parse_json(response.text or "") or {}

#                 # ==== عرض النتائج في Streamlit ====
#                 st.markdown("""<div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> 📊 حجم السوق (تقديري)</h3>
#                 </div>""", unsafe_allow_html=True)
#                 st.write(str(data.get("MarketSize", "-")))

#                 st.markdown("""<div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> 📈 معدل النمو السنوي (CAGR)</h3>
#                 </div>""", unsafe_allow_html=True)
#                 gr = data.get("GrowthRate", "-")
#                 gr_txt = f"{_to_float(gr):.2f}%"
#                 st.write(gr_txt)

#                 st.markdown(""" <div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> أقوى المنافسين في السعودية</h3>
#                 </div>""", unsafe_allow_html=True)
#                 comp = data.get("TopCompetitors", []) or []
#                 if comp:
#                     for c in comp:
#                         st.markdown(f"- {c}")
#                 else:
#                     st.warning("لم يتم العثور على منافسين.")

#                 st.markdown("""<div style="position: relative; background: rgba(0, 0, 0, 0.5); backdrop-filter: blur(8px); padding: 10px; border-radius: 8px;">
#                 <h3 style="color: white; text-align: center;"> تحليل SWOT</h3>
#                 </div>""", unsafe_allow_html=True)
#                 sw = data.get("SWOT", {}) or {}

#                 # Strengths, Weaknesses, Opportunities, Threats display logic...

#                 # تنزيل Word خارج النموذج
#                 filename = f"Market_Report_{store_market}{btype}_{country}.docx"
#                 export_market_report_to_docx(data, filename)
                
#                 # وضع زر التحميل بعد تقديم النموذج
#                 with open(filename, "rb") as f:
#                     st.download_button(
#                         "⬇ تنزيل تقرير السوق (Word)",
#                         data=f,
#                         file_name=filename,
#                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                     )

#             except Exception as e:
#                 st.error(f"❌ خطأ في تحليل السوق: {e}")

with tab5:
    with st.form("form_market_analysis"):
        st.subheader("📈 تحليل السوق السعودي ")

        # صف 1: مدخلان
        r1c1, r1c2 = st.columns(2)
        with r1c1:
            store_market = st.selectbox(
                "اختر المتجر",
                options=clients["StoreName"].unique(),
                index=None,
                placeholder="اختر المتجر",
                key="store_market",
            )
        with r1c2:
            btype = st.selectbox(
                "اختر نوع الفئة",
                options=["B2C", "B2B"],
                index=None,
                placeholder="اختر نوع الفئة",
                key="btype_market",
            )

        # صف 2: مدخلان
        r2c1, r2c2 = st.columns(2)
        with r2c1:
            country = st.selectbox(
                "اختر الدولة",
                options=countries,
                index=None,
                placeholder="اختر الدولة",
                key="country_market",
            )
        with r2c2:
            cities = df[df['الدولة'] == country]['المنطقة'].unique()
            cities_options = ["None"] + list(cities)
            cities_selected = st.multiselect("اختر المدن", cities_options, key="cities_market")

        if not cities_selected or "None" in cities_selected:
            selected_text = f"كل مدن {country}"
            cities_selected = list(cities)
        else:
            selected_text = ", ".join(cities_selected)

        # زر الإرسال داخل الفورم
        submitted = st.form_submit_button("تحليل السوق بالذكاء الاصطناعي")

    # بعد الفورم: تنفيذ التحليل وعرض النتائج
    if submitted:
        try:
            row_m = clients[clients["StoreName"] == store_market].iloc[0]
            category_market = row_m.get("SubCategory") if pd.notna(row_m.get("SubCategory")) else row_m["Category"]

            # ==== توليد التقرير ====
            prompt_market = f"""
            انت باحث تسويق متخصص في السعودية. 
            ✅ مسموح فقط باللغة العربية المبسطة.
            ❌ ممنوع استخدام أي كلمة أو جملة باللغة الإنجليزية.
            ✅ لو لازم تذكر مصطلحات عالمية، اكتبها بالعربية متبوعة بالاختصار بين أقواس، مثل:
            - معدل النمو السنوي المركب (CAGR)
            ✅ اجعل كل جزء من التقرير في شكل قائمة مرقمة (1. ... 2. ... 3. ...).
            ✅ كل نقطة لازم تكون جملة قصيرة ومباشرة (سطر واحد فقط).

            اعطني تقرير عن السوق السعودي في مجال "{category_market}" للفئة "{btype}" 
            في دولة {country} ومدن {selected_text}.

            يجب أن يتضمن التقرير:
            1. حجم السوق (بالريال السعودي أو عدد العملاء).
            2. معدل النمو السنوي المركب (CAGR).
            3. أقوى 3 منافسين حقيقيين.
            4. تحليل SWOT (نقاط القوة، الضعف، الفرص، التهديدات) – كل قسم مرقم.
            5. 3 توصيات عملية واضحة ومباشرة.

            النتيجة لازم تكون JSON فقط بالصيغة:
            {{
            "MarketSize": "...",
            "GrowthRate": 0.0,
            "TopCompetitors": [". ...", ". ...", ". ..."],
            "SWOT": {{"Strengths": [" ...", " ..."], "Weaknesses": [" ...", " ..."], "Opportunities": ["...", "..."], "Threats": [" ...", " ..."]}},
            "Recommendations": ["...", "...", "..."]
            }}
            """
            model = genai.GenerativeModel("models/gemini-2.5-flash")
            response = model.generate_content(prompt_market)
            data = _safe_parse_json(response.text or "") or {}

            # ==== دالة مساعدة لإنشاء div موحد لكل قسم ====
            def render_section(title, content_list, emoji=""):
                st.markdown(f"""
                <div style="
                    margin: 0;
                    padding: 10px;
                    position: relative;
                    background: rgba(0,0,0,0.5);
                    backdrop-filter: blur(8px);
                    border-radius: 8px;
                ">
                    <h3 style="color: white; text-align: center;">{emoji} {title}</h3>
                </div>
                """, unsafe_allow_html=True)
                for i, item in enumerate(content_list, 1):
                    st.markdown(f"""
                    <div style="
                        margin: 5px 0;
                        padding: 10px 14px;
                        position: relative;
                        background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
                        backdrop-filter: blur(10px);
                        border: 1px solid rgba(255,255,255,0.2);
                        box-shadow: 0 2px 8px rgba(0,0,0,0.3);
                        border-radius: 10px;
                    ">
                        <p style="
                            color: #f0f0f0;
                            font-size: 0.95rem;
                            font-family: 'Cairo', sans-serif;
                            line-height: 1.6;
                            text-align: right;
                            direction: rtl;
                        ">🔹 {i}. {item}</p>
                    </div>
                    """, unsafe_allow_html=True)

            # حجم السوق
            render_section("حجم السوق (تقديري)", [str(data.get("MarketSize","-"))], "📊")

            # معدل النمو السنوي
            gr = data.get("GrowthRate", "-")
            gr_txt = f"{_to_float(gr):.2f}%" if gr != "-" else "-"
            render_section("معدل النمو السنوي (CAGR)", [gr_txt], "📈")

            # أقوى المنافسين
            comp = data.get("TopCompetitors", []) or []
            if comp:
                render_section("أقوى المنافسين في السعودية", comp)
            else:
                st.warning("لم يتم العثور على منافسين.")

            # تحليل SWOT
            sw = data.get("SWOT", {}) or {}
            for title, key, emoji in [("نقاط القوة", "Strengths", "✅"), 
                                      ("نقاط الضعف", "Weaknesses", "⚠"),
                                      ("الفرص", "Opportunities", "💡"), 
                                      ("التهديدات", "Threats", "🚨")]:
                render_section(title, sw.get(key, []), emoji)

            # التوصيات
            recs = data.get("Recommendations", []) or []
            render_section("التوصيات", recs, "📌")

            # حفظ التقرير لزر التنزيل
            file_suffix = f"{country}{''.join(cities_selected) if cities_selected else 'عام'}"
            filename = f"Market_Report_{store_market}_{btype}_{file_suffix}.docx"
            export_market_report_to_docx(data, filename)

        except Exception as e:
            st.error(f"❌ خطأ في تحليل السوق: {e}")

    # زر التنزيل خارج الفورم
    if submitted:
        try:
            with open(filename, "rb") as f:
                st.download_button(
                    "⬇ تنزيل تقرير السوق (Word)",
                    data=f,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"❌ خطأ في تنزيل التقرير: {e}")

